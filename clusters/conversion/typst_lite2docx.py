#!/usr/bin/env python3
"""
Typst Lite to DOCX Generator

A lightweight Typst-inspired markup parser that generates DOCX files.
Supports a subset of Typst syntax for document formatting.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

class TypstLiteGenerator:
    def __init__(self):
        self.doc = Document()

        # Default: A4 paper met 2.5cm margins
        self.set_page_format('a4', 2.5)

        # Default font settings
        self.default_font = 'Century Gothic'
        self.default_size = 12

        # Numbered list counters per level
        # We houden een dictionary bij van level -> counter
        self.enum_counters = {}

        # Fix default Normal style - zet spacing op 0
        # Dit zorgt ervoor dat OpenOffice/LibreOffice ook 0 spacing toont
        try:
            normal_style = self.doc.styles['Normal']
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)
            normal_style.paragraph_format.line_spacing = 1.0  # Single line spacing
        except KeyError:
            # Als Normal style niet bestaat, geen probleem
            pass

    def set_page_format(self, paper='a4', margin=None):
        """Stel pagina formaat en marges in

        Args:
            paper: 'a4' of 'a5'
            margin: dict met margin waarden of float voor alle kanten
                   Kan bevatten: top, bottom, left, right, x, y, inside, outside
        """
        # Default margin
        if margin is None:
            margin = 2.5

        for section in self.doc.sections:
            # Stel papierformaat in
            if paper.lower() == 'a4':
                # A4: 210mm x 297mm
                section.page_width = Mm(210)
                section.page_height = Mm(297)
            elif paper.lower() == 'a5':
                # A5: 148mm x 210mm
                section.page_width = Mm(148)
                section.page_height = Mm(210)

            # Stel marges in
            if isinstance(margin, dict):
                # Individuele margins opgegeven
                # x en y zijn shortcuts
                x_margin = margin.get('x')
                y_margin = margin.get('y')

                # inside/outside voor alternating pages (booklet style)
                inside = margin.get('inside')
                outside = margin.get('outside')

                # Als inside/outside zijn gebruikt, map naar left/right
                # Note: echte alternating pages vereist meer complexe DOCX setup
                if inside is not None:
                    section.left_margin = Cm(inside)
                    section.right_margin = Cm(outside if outside is not None else inside)
                elif outside is not None:
                    section.right_margin = Cm(outside)
                    if inside is None:
                        section.left_margin = Cm(outside)

                # Individuele margins (overschrijven inside/outside als opgegeven)
                if margin.get('top') is not None:
                    section.top_margin = Cm(margin['top'])
                elif y_margin is not None:
                    section.top_margin = Cm(y_margin)

                if margin.get('bottom') is not None:
                    section.bottom_margin = Cm(margin['bottom'])
                elif y_margin is not None:
                    section.bottom_margin = Cm(y_margin)

                if margin.get('left') is not None:
                    section.left_margin = Cm(margin['left'])
                elif x_margin is not None and inside is None:
                    section.left_margin = Cm(x_margin)

                if margin.get('right') is not None:
                    section.right_margin = Cm(margin['right'])
                elif x_margin is not None and outside is None:
                    section.right_margin = Cm(x_margin)
            else:
                # Simpele float - alle kanten gelijk
                margin_cm = Cm(margin)
                section.top_margin = margin_cm
                section.bottom_margin = margin_cm
                section.left_margin = margin_cm
                section.right_margin = margin_cm

    def find_matching_bracket(self, text, start_pos):
        """Vind de bijbehorende sluitende bracket"""
        depth = 1
        for i in range(start_pos, len(text)):
            if text[i] == '[':
                depth += 1
            elif text[i] == ']':
                depth -= 1
                if depth == 0:
                    return i
        return -1

    def find_matching_paren(self, text, start_pos):
        """Vind de bijbehorende sluitende haakje"""
        depth = 1
        for i in range(start_pos, len(text)):
            if text[i] == '(':
                depth += 1
            elif text[i] == ')':
                depth -= 1
                if depth == 0:
                    return i
        return -1

    def parse_text_params(self, params_str):
        """Parse parameters zoals: size: 14pt, font: "Arial", fill: rgb("#FF0000")"""
        params = {
            'size': None,
            'font': None,
            'color': None
        }

        # Parse size: 14pt
        size_match = re.search(r'size:\s*(\d+)pt', params_str)
        if size_match:
            params['size'] = int(size_match.group(1))

        # Parse font: "FontName"
        font_match = re.search(r'font:\s*"([^"]+)"', params_str)
        if font_match:
            params['font'] = font_match.group(1)

        # Parse fill: rgb("#RRGGBB")
        color_match = re.search(r'fill:\s*rgb\(["\']#([0-9A-Fa-f]{6})["\']\)', params_str)
        if color_match:
            params['color'] = color_match.group(1)

        return params

    def parse_inline_markup(self, text, base_attrs):
        """Parse inline markup: *bold*, _italic_, #super[x], en #text()[content]"""
        parts = []
        i = 0

        while i < len(text):
            # Check voor #text(params)[content]
            if text[i:i+5] == '#text' and i + 5 < len(text) and text[i+5] == '(':
                paren_end = self.find_matching_paren(text, i + 6)
                if paren_end != -1 and paren_end + 1 < len(text) and text[paren_end + 1] == '[':
                    params_str = text[i+6:paren_end]
                    params = self.parse_text_params(params_str)

                    bracket_end = self.find_matching_bracket(text, paren_end + 2)
                    if bracket_end != -1:
                        content = text[paren_end + 2:bracket_end]

                        # Nieuwe attributes met de text params
                        new_attrs = base_attrs.copy()
                        if params['size']:
                            new_attrs['size'] = params['size']
                        if params['font']:
                            new_attrs['font'] = params['font']
                        if params['color']:
                            new_attrs['color'] = params['color']

                        # Recursief parsen van de content
                        nested_parts = self.parse_inline_markup(content, new_attrs)
                        parts.extend(nested_parts)

                        i = bracket_end + 1
                        continue

            # Check voor *bold*
            if text[i] == '*':
                # Zoek de sluitende *
                end = text.find('*', i + 1)
                if end != -1:
                    # Parse de content recursief
                    content = text[i+1:end]
                    new_attrs = base_attrs.copy()
                    new_attrs['bold'] = True
                    nested_parts = self.parse_inline_markup(content, new_attrs)
                    parts.extend(nested_parts)

                    i = end + 1
                    continue
                else:
                    # Geen sluitende *, behandel als gewone tekst
                    parts.append({
                        'text': '*',
                        **base_attrs
                    })
                    i += 1
                    continue

            # Check voor _italic_
            if text[i] == '_':
                # Zoek de sluitende _
                end = text.find('_', i + 1)
                if end != -1:
                    # Parse de content recursief
                    content = text[i+1:end]
                    new_attrs = base_attrs.copy()
                    new_attrs['italic'] = True
                    nested_parts = self.parse_inline_markup(content, new_attrs)
                    parts.extend(nested_parts)

                    i = end + 1
                    continue
                else:
                    # Geen sluitende _, behandel als gewone tekst
                    parts.append({
                        'text': '_',
                        **base_attrs
                    })
                    i += 1
                    continue

            # Check voor #smallcaps[text] of #sc[text]
            if (text[i:i+11] == '#smallcaps[' or text[i:i+4] == '#sc['):
                if text[i:i+11] == '#smallcaps[':
                    bracket_start = i + 11
                else:  # #sc[
                    bracket_start = i + 4

                bracket_end = self.find_matching_bracket(text, bracket_start)
                if bracket_end != -1:
                    content = text[bracket_start:bracket_end]
                    new_attrs = base_attrs.copy()
                    new_attrs['smallcaps'] = True
                    # Recursief parsen voor geneste markup
                    nested_parts = self.parse_inline_markup(content, new_attrs)
                    parts.extend(nested_parts)

                    i = bracket_end + 1
                    continue

            # Check voor #super[x]
            if text[i:i+6] == '#super' and i + 6 < len(text) and text[i+6] == '[':
                bracket_end = self.find_matching_bracket(text, i + 7)
                if bracket_end != -1:
                    content = text[i+7:bracket_end]
                    new_attrs = base_attrs.copy()
                    new_attrs['superscript'] = True
                    parts.append({
                        'text': content,
                        **new_attrs
                    })

                    i = bracket_end + 1
                    continue

            # Gewone tekst
            # Verzamel alle gewone tekst tot de volgende markup
            text_start = i
            while i < len(text) and text[i] not in ['*', '_', '#']:
                i += 1

            if i > text_start:
                parts.append({
                    'text': text[text_start:i],
                    **base_attrs
                })
            elif i == text_start:
                # Als we niet vooruit zijn gegaan, forceer vooruit
                i += 1

        # Als er geen markup was, return de hele tekst
        if not parts:
            parts = [{
                'text': text,
                **base_attrs
            }]

        return parts

    def parse_line_commands(self, line):
        """Parse Typst commando's zoals #align(), #text(), etc."""
        attrs = {
            'text': line,
            'size': self.default_size,
            'bold': False,
            'italic': False,
            'align': 'left',
            'font': self.default_font,
            'color': None,
            'superscript': False,
            'smallcaps': False
        }

        original_line = line

        # Parse #align(center)[...] of #align(right)[...]
        align_match = re.match(r'#align\((center|right)\)\[', line)
        if align_match:
            attrs['align'] = align_match.group(1)
            bracket_start = align_match.end()
            bracket_end = self.find_matching_bracket(line, bracket_start)
            if bracket_end != -1:
                line = line[bracket_start:bracket_end]

        # Parse #text(params)[...] - kan meerdere keren voorkomen (genest)
        while True:
            text_match = re.match(r'#text\(([^\)]+)\)\[', line)
            if not text_match:
                break

            params_str = text_match.group(1)
            params = self.parse_text_params(params_str)

            if params['size']:
                attrs['size'] = params['size']
            if params['font']:
                attrs['font'] = params['font']
            if params['color']:
                attrs['color'] = params['color']

            bracket_start = text_match.end()
            bracket_end = self.find_matching_bracket(line, bracket_start)
            if bracket_end != -1:
                line = line[bracket_start:bracket_end]
            else:
                break

        attrs['text'] = line
        return attrs

    def parse_margin_dict(self, margin_str):
        """Parse margin dictionary zoals: (top: 3cm, bottom: 2cm, x: 1.5cm)

        Ondersteunt zowel punten (1.5cm) als komma's (1,5cm) als decimaalteken.
        """
        margins = {}

        # Parse individuele margin waarden
        # Ondersteunt: top, bottom, left, right, x, y, inside, outside
        # Accepteert zowel punt als komma als decimaalteken
        patterns = {
            'top': r'top:\s*([0-9.,]+)cm',
            'bottom': r'bottom:\s*([0-9.,]+)cm',
            'left': r'left:\s*([0-9.,]+)cm',
            'right': r'right:\s*([0-9.,]+)cm',
            'x': r'x:\s*([0-9.,]+)cm',
            'y': r'y:\s*([0-9.,]+)cm',
            'inside': r'inside:\s*([0-9.,]+)cm',
            'outside': r'outside:\s*([0-9.,]+)cm',
        }

        for key, pattern in patterns.items():
            match = re.search(pattern, margin_str)
            if match:
                # Vervang komma door punt voor float parsing
                value_str = match.group(1).replace(',', '.')
                margins[key] = float(value_str)

        return margins if margins else None

    def parse_set_text(self, line):
        """Parse #set text() commando

        Ondersteunt:
        - #set text(font: "Times New Roman")
        - #set text(size: 14pt)
        - #set text(font: "Arial", size: 11pt)
        """
        # Extract parameters tussen haakjes
        match = re.search(r'#set\s+text\((.+)\)', line)
        if not match:
            return

        params_str = match.group(1)

        # Parse font parameter
        font_match = re.search(r'font:\s*"([^"]+)"', params_str)
        if font_match:
            self.default_font = font_match.group(1)

        # Parse size parameter
        size_match = re.search(r'size:\s*(\d+)pt', params_str)
        if size_match:
            self.default_size = int(size_match.group(1))

    def parse_set_page(self, line):
        """Parse #set page() commando

        Ondersteunt:
        - #set page(paper: "a4")
        - #set page(paper: "a5")
        - #set page(margin: 2.5cm)
        - #set page(paper: "a4", margin: 2.5cm)
        - #set page(margin: (top: 3cm, bottom: 2cm, x: 1.5cm))
        - #set page(margin: (inside: 2.5cm, outside: 2cm, y: 1.75cm))
        """
        # Extract parameters tussen haakjes
        match = re.search(r'#set\s+page\((.+)\)', line, re.DOTALL)
        if not match:
            return

        params_str = match.group(1)

        # Parse paper parameter
        paper_match = re.search(r'paper:\s*"(a4|a5)"', params_str)
        paper = paper_match.group(1) if paper_match else None

        # Parse margin parameter
        # Eerst checken of het een dict is: margin: (...)
        margin_dict_match = re.search(r'margin:\s*\(([^)]+)\)', params_str)
        if margin_dict_match:
            # Complexe margin met meerdere waarden
            margin_str = margin_dict_match.group(1)
            margin = self.parse_margin_dict(margin_str)
        else:
            # Simpele margin: margin: 2.5cm of margin: 2,5cm
            margin_simple_match = re.search(r'margin:\s*([0-9.,]+)cm', params_str)
            if margin_simple_match:
                # Vervang komma door punt voor float parsing
                value_str = margin_simple_match.group(1).replace(',', '.')
                margin = float(value_str)
            else:
                margin = None

        # Als we iets gevonden hebben, pas het toe
        if paper or margin:
            current_paper = paper if paper else 'a4'  # Default blijft a4
            current_margin = margin if margin is not None else 2.5  # Default blijft 2.5cm

            self.set_page_format(current_paper, current_margin)

    def parse_v_spacing(self, v_command):
        """Parse #v() commando en return spacing in punten

        Ondersteunt:
        - #v(12pt) - punten
        - #v(1em) - em (1em = 12pt)
        - #v(0.5em) - fractional em
        """
        # Extract de waarde tussen haakjes
        match = re.search(r'#v\(([0-9.]+)(pt|em)\)', v_command)
        if not match:
            # Default fallback naar 12pt (1em)
            return 12

        value = float(match.group(1))
        unit = match.group(2)

        if unit == 'pt':
            return value
        elif unit == 'em':
            # 1em = 12pt (standaard font size)
            return value * 12

        return 12  # Default

    def parse_list_item(self, line):
        """Parse een list item en return (is_list, indent_level, content)

        List items beginnen met - gevolgd door een spatie.
        Indentatie bepaalt het nesting level (2 spaties = 1 level).

        Returns:
            tuple: (is_list_item, indent_level, content_text)
        """
        # Check of de regel start met optionele spaties gevolgd door - en een spatie
        match = re.match(r'^( *)- (.+)$', line)
        if not match:
            return (False, 0, '')

        indent_spaces = len(match.group(1))
        content = match.group(2)

        # Bereken indent level: elke 2 spaties is 1 level
        indent_level = indent_spaces // 2

        return (True, indent_level, content)

    def add_list_item(self, content, level):
        """Voeg een list item toe met de juiste indent level

        Args:
            content: De tekst van het list item
            level: Het indentatie level (0 = top level, 1 = nested, etc.)
        """
        # Parse de content voor inline markup
        base_attrs = {
            'bold': False,
            'italic': False,
            'color': None,
            'font': self.default_font,
            'size': self.default_size,
            'superscript': False,
            'smallcaps': False
        }
        parts = self.parse_inline_markup(content, base_attrs)

        # Maak de paragraph met list bullet style
        p = self.doc.add_paragraph()

        # Stel indentatie in gebaseerd op level
        # Level 0: standaard bullet •
        # Level 1: nested bullet ‣ (via indent)
        # Level 2: nested bullet – (via meer indent)
        indent_per_level = Cm(0.5)  # 0.5cm per level
        bullet_markers = ['•', '‣', '–']  # Typst default markers

        # Bepaal de bullet marker op basis van level
        marker = bullet_markers[level % len(bullet_markers)]

        # Stel paragraph format in
        p.paragraph_format.left_indent = indent_per_level * (level + 1)
        p.paragraph_format.first_line_indent = Cm(-0.5)  # Hanging indent voor bullet
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Voeg de bullet marker toe
        marker_run = p.add_run(f'{marker} ')
        marker_run.font.name = self.default_font
        marker_run.font.size = Pt(self.default_size)

        # Voeg de content runs toe
        for part in parts:
            run = p.add_run(part['text'])
            run.font.size = Pt(part['size'])
            run.font.bold = part['bold']
            run.font.italic = part['italic']
            run.font.name = part['font']
            if part.get('superscript'):
                run.font.superscript = True
            if part.get('smallcaps'):
                run.font.small_caps = True
            if part['color']:
                run.font.color.rgb = RGBColor(
                    int(part['color'][0:2], 16),
                    int(part['color'][2:4], 16),
                    int(part['color'][4:6], 16)
                )

        return p

    def parse_enum_item(self, line):
        """Parse een numbered list item en return (is_enum, indent_level, content)

        Enum items beginnen met + gevolgd door een spatie.
        Indentatie bepaalt het nesting level (2 spaties = 1 level).

        Returns:
            tuple: (is_enum_item, indent_level, content_text)
        """
        # Check of de regel start met optionele spaties gevolgd door + en een spatie
        match = re.match(r'^( *)\+ (.+)$', line)
        if not match:
            return (False, 0, '')

        indent_spaces = len(match.group(1))
        content = match.group(2)

        # Bereken indent level: elke 2 spaties is 1 level
        indent_level = indent_spaces // 2

        return (True, indent_level, content)

    def get_enum_number_format(self, level, number):
        """Converteer een nummer naar het juiste format voor het gegeven level

        Args:
            level: Het nesting level (0, 1, 2, ...)
            number: Het huidige nummer (1, 2, 3, ...)

        Returns:
            str: Het geformatteerde nummer (bijv. "1.", "a.", "i.")
        """
        # Verschillende numbering formats per level
        # Level 0: 1. 2. 3.
        # Level 1: a. b. c.
        # Level 2: i. ii. iii.
        # Level 3+: cyclus herhaalt

        level_mod = level % 3

        if level_mod == 0:
            # Arabische cijfers: 1. 2. 3.
            return f"{number}."
        elif level_mod == 1:
            # Kleine letters: a. b. c.
            # a = 97 in ASCII
            if number <= 26:
                return f"{chr(96 + number)}."
            else:
                # Voor getallen > 26, gebruik aa, ab, etc.
                return f"{chr(96 + ((number - 1) // 26))}{chr(97 + ((number - 1) % 26))}."
        else:  # level_mod == 2
            # Romeinse cijfers: i. ii. iii.
            roman_numerals = [
                '', 'i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x',
                'xi', 'xii', 'xiii', 'xiv', 'xv', 'xvi', 'xvii', 'xviii', 'xix', 'xx'
            ]
            if number < len(roman_numerals):
                return f"{roman_numerals[number]}."
            else:
                # Fallback voor grote getallen
                return f"{number}."

    def reset_deeper_enum_counters(self, level):
        """Reset alle enum counters die dieper zijn dan het gegeven level

        Dit zorgt ervoor dat wanneer we terugkeren naar een hoger level,
        de geneste nummering opnieuw begint.
        """
        levels_to_remove = [l for l in self.enum_counters.keys() if l > level]
        for l in levels_to_remove:
            del self.enum_counters[l]

    def add_enum_item(self, content, level):
        """Voeg een numbered list item toe met de juiste nummering

        Args:
            content: De tekst van het enum item
            level: Het indentatie level (0 = top level, 1 = nested, etc.)
        """
        # Reset diepere levels als we terug zijn naar een hoger level
        self.reset_deeper_enum_counters(level)

        # Increment de counter voor dit level (of start bij 1)
        if level not in self.enum_counters:
            self.enum_counters[level] = 1
        else:
            self.enum_counters[level] += 1

        current_number = self.enum_counters[level]

        # Parse de content voor inline markup
        base_attrs = {
            'bold': False,
            'italic': False,
            'color': None,
            'font': self.default_font,
            'size': self.default_size,
            'superscript': False,
            'smallcaps': False
        }
        parts = self.parse_inline_markup(content, base_attrs)

        # Maak de paragraph met numbered list style
        p = self.doc.add_paragraph()

        # Stel indentatie in gebaseerd op level
        indent_per_level = Cm(0.5)  # 0.5cm per level

        # Bepaal het nummer format op basis van level
        number_text = self.get_enum_number_format(level, current_number)

        # Stel paragraph format in
        p.paragraph_format.left_indent = indent_per_level * (level + 1)
        p.paragraph_format.first_line_indent = Cm(-0.5)  # Hanging indent voor nummer
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Voeg het nummer toe
        number_run = p.add_run(f'{number_text} ')
        number_run.font.name = self.default_font
        number_run.font.size = Pt(self.default_size)

        # Voeg de content runs toe
        for part in parts:
            run = p.add_run(part['text'])
            run.font.size = Pt(part['size'])
            run.font.bold = part['bold']
            run.font.italic = part['italic']
            run.font.name = part['font']
            if part.get('superscript'):
                run.font.superscript = True
            if part.get('smallcaps'):
                run.font.small_caps = True
            if part['color']:
                run.font.color.rgb = RGBColor(
                    int(part['color'][0:2], 16),
                    int(part['color'][2:4], 16),
                    int(part['color'][4:6], 16)
                )

        return p

    def add_line(self, attrs):
        """Voeg een regel toe met attributen"""
        p = self.doc.add_paragraph()

        # Alignment
        if attrs['align'] == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif attrs['align'] == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Spacing - zet alles op 0 voor volledige controle
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Parse inline markup
        base_attrs = {
            'bold': attrs['bold'],
            'italic': attrs['italic'],
            'color': attrs['color'],
            'font': attrs['font'],
            'size': attrs['size'],
            'superscript': attrs['superscript'],
            'smallcaps': attrs['smallcaps']
        }
        parts = self.parse_inline_markup(attrs['text'], base_attrs)

        # Voeg elke part toe als een run
        for part in parts:
            run = p.add_run(part['text'])
            run.font.size = Pt(part['size'])
            run.font.bold = part['bold']
            run.font.italic = part['italic']
            run.font.name = part['font']
            if part.get('superscript'):
                run.font.superscript = True
            if part.get('smallcaps'):
                run.font.small_caps = True
            if part['color']:
                run.font.color.rgb = RGBColor(
                    int(part['color'][0:2], 16),
                    int(part['color'][2:4], 16),
                    int(part['color'][4:6], 16)
                )

        return p

    def process_file(self, input_file):
        """Verwerk het markup bestand"""
        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            line = line.rstrip()

            # #set page() - pagina setup
            if line.strip().startswith('#set page('):
                self.parse_set_page(line.strip())
                continue

            # #set text() - font/size setup
            if line.strip().startswith('#set text('):
                self.parse_set_text(line.strip())
                continue

            # #v(12pt) of #v(1em) - lege regel met specifieke font size
            if line.strip().startswith('#v('):
                spacing_pt = self.parse_v_spacing(line.strip())
                # Maak een lege paragraph met de juiste font properties
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                # Voeg een spatie toe met de juiste font size en lettertype
                # Een spatie is nodig zodat de fontgrootte effect heeft
                run = p.add_run(' ')
                run.font.size = Pt(spacing_pt)
                run.font.name = self.default_font
                continue

            # #pagebreak()
            if line.strip() == '#pagebreak()':
                self.doc.add_page_break()
                continue

            # Skip volledig lege regels
            if not line.strip():
                continue

            # Check voor bullet list items (- Item)
            is_list, level, content = self.parse_list_item(line)
            if is_list:
                self.add_list_item(content, level)
                continue

            # Check voor numbered list items (+ Item)
            is_enum, level, content = self.parse_enum_item(line)
            if is_enum:
                self.add_enum_item(content, level)
                continue

            # Parse en voeg toe
            attrs = self.parse_line_commands(line)
            if attrs['text']:
                self.add_line(attrs)
            else:
                self.doc.add_paragraph()

    def save(self, output_file):
        """Save het document"""
        self.doc.save(output_file)
        print(f"DOCX gegenereerd: {output_file}")

def main():
    import sys

    if len(sys.argv) < 2:
        print("Gebruik: python typst_lite2docx.py <input.txt> [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'output.docx'

    generator = TypstLiteGenerator()
    generator.process_file(input_file)
    generator.save(output_file)

if __name__ == '__main__':
    main()
